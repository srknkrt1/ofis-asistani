from faster_whisper import WhisperModel
from docx import Document
import os
import tempfile
from pydub import AudioSegment
AudioSegment.converter = "/usr/bin/ffmpeg"

# Modeli sadece bir kez yükle (CPU için optimize)
model = WhisperModel("small", compute_type="int8")  # "base" veya "tiny" ile değiştirebilirsin

def transkripte_cevir(dosya_adi, dil="tr"):
    segments, info = model.transcribe(dosya_adi, language=dil)
    metin = ""
    for segment in segments:
        metin += segment.text.strip() + " "
    return metin.strip()

def transkript_docx_yaz(metin, dosya_adi="transkript.docx"):
    doc = Document()
    doc.add_heading("Transkript", level=1)
    doc.add_paragraph(metin)

    temp_path = os.path.join(tempfile.gettempdir(), dosya_adi)
    doc.save(temp_path)
    return temp_path
