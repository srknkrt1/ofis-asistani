
from docx import Document
import os
import tempfile

def transkripte_cevir(dosya_adi, dil="tr"):
    model = whisper.load_model("small")
    sonuc = model.transcribe(dosya_adi, language=dil)
    return sonuc["text"]

def transkript_docx_yaz(metin, dosya_adi="transkript.docx"):
    doc = Document()
    doc.add_heading("Transkript", level=1)
    doc.add_paragraph(metin)

    temp_path = os.path.join(tempfile.gettempdir(), dosya_adi)
    doc.save(temp_path)
    return temp_path

