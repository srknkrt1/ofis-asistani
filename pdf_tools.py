from PyPDF2 import PdfMerger, PdfReader, PdfWriter
import subprocess
from pdf2docx import Converter
from docx2pdf import convert
import pytesseract
from pdf2image import convert_from_path
from docx import Document

def birlestir_pdf_listesi(pdf_filepaths, cikti_yolu):
    merger = PdfMerger()
    for path in pdf_filepaths:
        merger.append(path)
    merger.write(cikti_yolu)
    merger.close()

def bol_pdf(giris_pdf, bas_sayfa, bit_sayfa, cikti_pdf):
    reader = PdfReader(giris_pdf)
    writer = PdfWriter()
    for i in range(bas_sayfa - 1, bit_sayfa):
        writer.add_page(reader.pages[i])
    with open(cikti_pdf, "wb") as f:
        writer.write(f)

def dondur_pdf(giris_pdf, sayfa_no, aci, cikti_pdf):
    reader = PdfReader(giris_pdf)
    writer = PdfWriter()
    for i, sayfa in enumerate(reader.pages):
        if i == sayfa_no - 1:
            sayfa.rotate(aci)
        writer.add_page(sayfa)
    with open(cikti_pdf, "wb") as f:
        writer.write(f)

def sikistir_pdf(giris_pdf, cikti_pdf, kalite="/ebook"):
    gs_path = "/usr/bin/gs"
    subprocess.call([
        gs_path, "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.4",
        f"-dPDFSETTINGS={kalite}", "-dNOPAUSE", "-dQUIET", "-dBATCH",
        f"-sOutputFile={cikti_pdf}", giris_pdf
    ])

def pdf_to_word(giris_pdf, cikti_docx):
    cv = Converter(giris_pdf)
    cv.convert(cikti_docx)
    cv.close()

def word_to_pdf(giris_docx, cikti_pdf):
    convert(giris_docx, cikti_pdf)

def pdf_ocr_to_word(pdf_path, output_docx):
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"
    pages = convert_from_path(pdf_path)
    doc = Document()
    for i, page in enumerate(pages, 1):
        text = pytesseract.image_to_string(page, lang="tur+eng")
        doc.add_paragraph(f"--- Sayfa {i} ---")
        doc.add_paragraph(text)
        doc.add_page_break()
    doc.save(output_docx)
